"""
RAG Document Loader
Extracts text from PDF, DOCX, and TXT files with page-level preservation.
Reuses proven multi-format extraction from Project 2.
"""

import os
import io
from typing import Dict, Any, List, Optional, Union, BinaryIO


class ExtractedDocument:
    def __init__(
        self,
        filename: str,
        file_type: str,
        full_text: str = "",
        pages: Optional[List[Dict[str, Any]]] = None,
        metadata: Optional[Dict[str, Any]] = None,
        error: Optional[str] = None,
        is_valid: bool = True,
    ):
        self.filename = filename
        self.file_type = file_type.lower().lstrip(".")
        self.full_text = full_text or ""
        self.pages = pages or []
        self.metadata = metadata or {}
        self.error = error
        self.is_valid = is_valid and (error is None) and (len(self.full_text.strip()) > 0)
        
        clean = self.full_text.strip()
        self.total_chars = len(clean)
        self.total_words = len(clean.split()) if clean else 0
        self.page_count = len(self.pages) if self.pages else (1 if self.total_chars > 0 else 0)


def extract_text_from_pdf(file_source: Union[str, BinaryIO, bytes], filename: Optional[str] = None) -> ExtractedDocument:
    doc_name = filename or "document.pdf"
    try:
        import pypdf
    except ImportError:
        return ExtractedDocument(doc_name, "pdf", error="pypdf is not installed", is_valid=False)

    pages_data = []
    full_text_parts = []

    try:
        stream = io.BytesIO(file_source) if isinstance(file_source, (bytes, bytearray)) else (open(file_source, "rb") if isinstance(file_source, str) else file_source)
        reader = pypdf.PdfReader(stream)
        
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            pages_data.append({"page_number": idx, "text": text, "char_count": len(text)})
            if text:
                full_text_parts.append(f"--- [Page {idx}] ---\n{text}")

        full_text = "\n\n".join(full_text_parts).strip()
        return ExtractedDocument(doc_name, "pdf", full_text=full_text, pages=pages_data, is_valid=bool(full_text))
    except Exception as e:
        return ExtractedDocument(doc_name, "pdf", error=str(e), is_valid=False)


def extract_text_from_docx(file_source: Union[str, BinaryIO, bytes], filename: Optional[str] = None) -> ExtractedDocument:
    doc_name = filename or "document.docx"
    try:
        import docx
    except ImportError:
        return ExtractedDocument(doc_name, "docx", error="python-docx is not installed", is_valid=False)

    try:
        stream = io.BytesIO(file_source) if isinstance(file_source, (bytes, bytearray)) else (open(file_source, "rb") if isinstance(file_source, str) else file_source)
        doc = docx.Document(stream)
        sections = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(sections).strip()
        return ExtractedDocument(doc_name, "docx", full_text=full_text, pages=[{"page_number": 1, "text": full_text}], is_valid=bool(full_text))
    except Exception as e:
        return ExtractedDocument(doc_name, "docx", error=str(e), is_valid=False)


def extract_text_from_txt(file_source: Union[str, BinaryIO, bytes], filename: Optional[str] = None) -> ExtractedDocument:
    doc_name = filename or "document.txt"
    try:
        raw_bytes = bytes(file_source) if isinstance(file_source, (bytes, bytearray)) else (open(file_source, "rb").read() if isinstance(file_source, str) else file_source.read())
        if isinstance(raw_bytes, str):
            raw_bytes = raw_bytes.encode("utf-8")

        text = None
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = raw_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            text = raw_bytes.decode("utf-8", errors="ignore")

        clean = text.strip()
        return ExtractedDocument(doc_name, "txt", full_text=clean, pages=[{"page_number": 1, "text": clean}], is_valid=bool(clean))
    except Exception as e:
        return ExtractedDocument(doc_name, "txt", error=str(e), is_valid=False)


def extract_text_from_csv(file_source: Union[str, BinaryIO, bytes], filename: Optional[str] = None) -> ExtractedDocument:
    doc_name = filename or "document.csv"
    try:
        raw_bytes = bytes(file_source) if isinstance(file_source, (bytes, bytearray)) else (open(file_source, "rb").read() if isinstance(file_source, str) else file_source.read())
        if isinstance(raw_bytes, str):
            raw_bytes = raw_bytes.encode("utf-8")

        text_content = None
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text_content = raw_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if text_content is None:
            text_content = raw_bytes.decode("utf-8", errors="ignore")

        import csv
        reader = csv.reader(io.StringIO(text_content))
        rows = list(reader)
        if not rows:
            return ExtractedDocument(doc_name, "csv", full_text="", pages=[], is_valid=False)

        header = [h.strip() for h in rows[0]]
        formatted_rows = []
        for row_idx, row in enumerate(rows[1:], start=1):
            row_items = []
            for col_idx, val in enumerate(row):
                col_name = header[col_idx] if col_idx < len(header) and header[col_idx] else f"Col_{col_idx+1}"
                row_items.append(f"{col_name}: {val.strip()}")
            if row_items:
                formatted_rows.append(f"Row {row_idx}: " + ", ".join(row_items))

        full_text = f"Dataset: {doc_name}\nColumns: {', '.join(header)}\n\n" + "\n".join(formatted_rows)
        return ExtractedDocument(
            doc_name,
            "csv",
            full_text=full_text,
            pages=[{"page_number": 1, "text": full_text}],
            metadata={"row_count": len(rows) - 1, "col_count": len(header)},
            is_valid=bool(full_text.strip()),
        )
    except Exception as e:
        return ExtractedDocument(doc_name, "csv", error=str(e), is_valid=False)


def extract_text_from_json(file_source: Union[str, BinaryIO, bytes], filename: Optional[str] = None) -> ExtractedDocument:
    doc_name = filename or "document.json"
    try:
        raw_bytes = bytes(file_source) if isinstance(file_source, (bytes, bytearray)) else (open(file_source, "rb").read() if isinstance(file_source, str) else file_source.read())
        if isinstance(raw_bytes, str):
            raw_bytes = raw_bytes.encode("utf-8")

        text_content = None
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text_content = raw_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if text_content is None:
            text_content = raw_bytes.decode("utf-8", errors="ignore")

        import json
        data = json.loads(text_content)

        def format_json_node(obj, depth=0) -> str:
            indent = "  " * depth
            if isinstance(obj, dict):
                lines = []
                for k, v in obj.items():
                    if isinstance(v, (dict, list)):
                        lines.append(f"{indent}{k}:\n{format_json_node(v, depth+1)}")
                    else:
                        lines.append(f"{indent}{k}: {v}")
                return "\n".join(lines)
            elif isinstance(obj, list):
                lines = []
                for i, item in enumerate(obj, 1):
                    if isinstance(item, (dict, list)):
                        lines.append(f"{indent}[Item {i}]:\n{format_json_node(item, depth+1)}")
                    else:
                        lines.append(f"{indent}- {item}")
                return "\n".join(lines)
            else:
                return f"{indent}{str(obj)}"

        formatted_text = f"JSON Document: {doc_name}\n\n" + format_json_node(data)
        return ExtractedDocument(
            doc_name,
            "json",
            full_text=formatted_text,
            pages=[{"page_number": 1, "text": formatted_text}],
            is_valid=bool(formatted_text.strip()),
        )
    except Exception as e:
        return ExtractedDocument(doc_name, "json", error=str(e), is_valid=False)


def load_document(file_source: Union[str, BinaryIO, bytes], filename: str) -> ExtractedDocument:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_source, filename)
    elif ext == ".docx":
        return extract_text_from_docx(file_source, filename)
    elif ext == ".csv":
        return extract_text_from_csv(file_source, filename)
    elif ext == ".json":
        return extract_text_from_json(file_source, filename)
    else:
        # Defaults to plain text / markdown
        return extract_text_from_txt(file_source, filename)

