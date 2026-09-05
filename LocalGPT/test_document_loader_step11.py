"""
LocalGPT: Step 11 Document Upload Verification Test Suite
Tests:
1. TXT Extraction: Normal Python reading across UTF-8, multi-paragraph text.
2. DOCX Extraction: python-docx parsing for paragraphs and tables.
3. PDF Extraction: pypdf parsing with page preservation.
4. Empty & Corrupted Files: Graceful error handling without crashes.
5. Directory Storage & Management: File persistence and listing in data/documents.
"""

import os
import sys
import io
import time
import shutil
import tempfile

# Ensure UTF-8 stdout encoding for Windows console
if sys.platform == "win32" and hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# Ensure LocalGPT path is in sys.path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_loader import (
    extract_text_from_txt,
    extract_text_from_docx,
    extract_text_from_pdf,
    load_and_extract_document,
    save_uploaded_file,
    list_saved_documents,
    delete_saved_document,
    get_documents_directory,
    ExtractedDocument,
)

import docx
from pypdf import PdfWriter


def create_sample_pdf(file_path: str):
    """
    Creates a valid 2-page test PDF file using raw PDF stream syntax.
    """
    # Minimal valid PDF with 2 pages and stream text
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 5 0 R /Resources << /Font << /F1 7 0 R >> >> >> endobj\n"
        b"4 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 6 0 R /Resources << /Font << /F1 7 0 R >> >> >> endobj\n"
        b"5 0 obj << /Length 55 >> stream\n"
        b"BT /F1 12 Tf 72 712 Td (LocalGPT Step 11 PDF Page One Content) Tj ET\n"
        b"endstream endobj\n"
        b"6 0 obj << /Length 55 >> stream\n"
        b"BT /F1 12 Tf 72 712 Td (LocalGPT Step 11 PDF Page Two Content) Tj ET\n"
        b"endstream endobj\n"
        b"7 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n"
        b"0 8\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000236 00000 n \n"
        b"0000000357 00000 n \n"
        b"0000000463 00000 n \n"
        b"0000000569 00000 n \n"
        b"trailer << /Size 8 /Root 1 0 R >>\n"
        b"startxref\n"
        b"647\n"
        b"%%EOF\n"
    )
    with open(file_path, "wb") as f:
        f.write(pdf_content)


def create_sample_docx(file_path: str):
    """
    Creates a valid sample .docx document using python-docx with headings, paragraphs, and a table.
    """
    doc = docx.Document()
    doc.add_heading("LocalGPT Architecture Report", level=1)
    doc.add_paragraph("LocalGPT is an offline, private conversational AI designed for local edge inference.")
    doc.add_paragraph("Step 11 introduces comprehensive multi-format document loading and parsing.")
    
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Format"
    hdr_cells[1].text = "Library"
    hdr_cells[2].text = "Support Status"
    
    row_cells = table.add_row().cells
    row_cells[0].text = "PDF"
    row_cells[1].text = "pypdf"
    row_cells[2].text = "Full Page Preservation"
    
    doc.save(file_path)


def create_sample_txt(file_path: str):
    """
    Creates a sample .txt document.
    """
    text_content = (
        "LocalGPT Project Overview\n\n"
        "This is a sample plain text document used to verify Step 11 text extraction.\n"
        "It contains multiple paragraphs and UTF-8 symbols: α, β, γ, 🤖, ⚡.\n"
        "Plain text files are parsed natively using Python standard file reading with automatic encoding fallback."
    )
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text_content)


def run_step11_tests():
    print("=" * 75)
    print("LocalGPT Step 11: Document Upload & Extraction Verification Suite")
    print("=" * 75)

    test_dir = tempfile.mkdtemp(prefix="localgpt_test_docs_")
    results = {}

    try:
        # File paths
        txt_path = os.path.join(test_dir, "sample_test.txt")
        docx_path = os.path.join(test_dir, "sample_report.docx")
        pdf_path = os.path.join(test_dir, "sample_doc.pdf")
        empty_path = os.path.join(test_dir, "empty.txt")
        corrupted_pdf_path = os.path.join(test_dir, "corrupted.pdf")

        # 1. Create sample test files
        create_sample_txt(txt_path)
        create_sample_docx(docx_path)
        create_sample_pdf(pdf_path)
        with open(empty_path, "w") as f:
            f.write("")
        with open(corrupted_pdf_path, "wb") as f:
            f.write(b"NOT_A_REAL_PDF_FILE_HEADER_GARBAGE")

        # -------------------------------------------------------------
        # Test 1: Plain Text (TXT) Extraction
        # -------------------------------------------------------------
        print("\n[Test 1/5] Testing Plain Text (TXT) Extraction...")
        txt_doc = extract_text_from_txt(txt_path)
        print(f"  [+] Extracted: '{txt_doc.filename}' ({txt_doc.total_words} words, {txt_doc.total_chars} chars)")
        print(f"  [+] Excerpt: {txt_doc.full_text[:80]}...")
        assert txt_doc.is_valid, "TXT doc should be valid"
        assert "LocalGPT Project Overview" in txt_doc.full_text
        assert "α, β, γ" in txt_doc.full_text or "🤖" in txt_doc.full_text
        assert txt_doc.total_words > 20
        results["Plain Text (TXT) Extraction"] = True

        # -------------------------------------------------------------
        # Test 2: Word Document (DOCX) Extraction
        # -------------------------------------------------------------
        print("\n[Test 2/5] Testing Word Document (DOCX) Extraction...")
        docx_doc = extract_text_from_docx(docx_path)
        print(f"  [+] Extracted: '{docx_doc.filename}' ({docx_doc.total_words} words, {docx_doc.total_chars} chars)")
        print(f"  [+] Excerpt: {docx_doc.full_text[:100]}...")
        assert docx_doc.is_valid, "DOCX doc should be valid"
        assert "LocalGPT Architecture Report" in docx_doc.full_text
        assert "pypdf" in docx_doc.full_text
        assert docx_doc.metadata.get("table_count", 0) >= 1
        results["Word Document (DOCX) Extraction"] = True

        # -------------------------------------------------------------
        # Test 3: PDF Document (PDF) Extraction & Page Preservation
        # -------------------------------------------------------------
        print("\n[Test 3/5] Testing PDF Extraction & Page Preservation (pypdf)...")
        pdf_doc = extract_text_from_pdf(pdf_path)
        print(f"  [+] Extracted: '{pdf_doc.filename}' ({pdf_doc.page_count} pages, {pdf_doc.total_words} words)")
        for page in pdf_doc.pages:
            print(f"      - Page {page['page_number']}: '{page['text']}' ({page['word_count']} words)")
        assert pdf_doc.is_valid, "PDF doc should be valid"
        assert pdf_doc.page_count == 2, f"Expected 2 pages, got {pdf_doc.page_count}"
        assert "Page One" in pdf_doc.pages[0]["text"]
        assert "Page Two" in pdf_doc.pages[1]["text"]
        results["PDF Extraction & Page Preservation"] = True

        # -------------------------------------------------------------
        # Test 4: Empty & Corrupted Files Error Handling
        # -------------------------------------------------------------
        print("\n[Test 4/5] Testing Empty & Corrupted Files Safe Handling...")
        empty_res = extract_text_from_txt(empty_path)
        print(f"  [+] Empty File Result: is_valid={empty_res.is_valid}, error='{empty_res.error}'")
        assert not empty_res.is_valid
        assert empty_res.error is not None

        corrupted_res = extract_text_from_pdf(corrupted_pdf_path)
        print(f"  [+] Corrupted PDF Result: is_valid={corrupted_res.is_valid}, error='{corrupted_res.error}'")
        assert not corrupted_res.is_valid
        assert corrupted_res.error is not None
        results["Empty & Corrupted Files Handling"] = True

        # -------------------------------------------------------------
        # Test 5: Document Persistence & Storage in data/documents
        # -------------------------------------------------------------
        print("\n[Test 5/5] Testing Storage Persistence & Directory Management...")
        docs_dir = get_documents_directory()
        
        # Save sample files into data/documents
        with open(txt_path, "rb") as f:
            saved_txt = save_uploaded_file(f.read(), "step11_test_sample.txt")
        with open(docx_path, "rb") as f:
            saved_docx = save_uploaded_file(f.read(), "step11_test_sample.docx")
        with open(pdf_path, "rb") as f:
            saved_pdf = save_uploaded_file(f.read(), "step11_test_sample.pdf")

        # Verify listed documents
        saved_list = list_saved_documents()
        filenames = [d["filename"] for d in saved_list]
        print(f"  [+] Saved documents found in data/documents: {filenames}")
        assert "step11_test_sample.txt" in filenames
        assert "step11_test_sample.docx" in filenames
        assert "step11_test_sample.pdf" in filenames

        # Delete test artifacts from data/documents
        del1 = delete_saved_document("step11_test_sample.txt")
        del2 = delete_saved_document("step11_test_sample.docx")
        del3 = delete_saved_document("step11_test_sample.pdf")
        assert del1 and del2 and del3
        print("  [+] Cleanup verified: test documents safely removed.")
        results["Storage Persistence & Management"] = True

    finally:
        shutil.rmtree(test_dir, ignore_errors=True)

    # Summary Report
    print("\n" + "=" * 75)
    print("STEP 11 DOCUMENT UPLOAD VERIFICATION RESULTS:")
    print("=" * 75)
    all_passed = True
    for test_name, passed in results.items():
        status_box = "[X]" if passed else "[ ]"
        status_str = "PASS" if passed else "FAIL"
        print(f"  {status_box} {test_name:<48} : {status_str}")
        if not passed:
            all_passed = False

    print("=" * 75)
    if all_passed:
        print("ALL STEP 11 DOCUMENT UPLOAD TESTS PASSED SUCCESSFULLY! (5/5)")
    else:
        print("SOME TESTS FAILED.")
    print("=" * 75)
    return all_passed


if __name__ == "__main__":
    success = run_step11_tests()
    if not success:
        sys.exit(1)
