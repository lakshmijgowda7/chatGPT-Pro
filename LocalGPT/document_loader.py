"""
LocalGPT: Document Loader (Step 11: Document Upload)
Provides robust extraction, validation, and storage for:
- PDF (via pypdf with page preservation)
- DOCX (via python-docx with paragraph & table extraction)
- TXT / Markdown (via native Python text reading with multi-encoding support)

Handles empty or corrupted files gracefully and manages storage in data/documents/.
"""

import os
import io
import time
import shutil
from typing import List, Dict, Any, Optional, Union, BinaryIO

# -------------------------------------------------------------
# EXTRACTED DOCUMENT DATA STRUCTURE
# -------------------------------------------------------------
class ExtractedDocument:
    """
    Represents an extracted document with text, page breakdown, and metadata.
    """
    def __init__(
        self,
        filename: str,
        file_type: str,
        full_text: str = "",
        pages: Optional[List[Dict[str, Any]]] = None,
        metadata: Optional[Dict[str, Any]] = None,
        error: Optional[str] = None,
        is_valid: bool = True,
    ):
        self.filename = filename
        self.file_type = file_type.lower().lstrip(".")
        self.full_text = full_text or ""
        self.pages = pages or []
        self.metadata = metadata or {}
        self.error = error
        self.is_valid = is_valid and (error is None) and (len(self.full_text.strip()) > 0)
        
        # Calculate text statistics
        text_clean = self.full_text.strip()
        words = text_clean.split() if text_clean else []
        self.total_chars = len(text_clean)
        self.total_words = len(words)
        self.page_count = len(self.pages) if self.pages else (1 if self.total_chars > 0 else 0)

        # Update metadata stats
        self.metadata["total_chars"] = self.total_chars
        self.metadata["total_words"] = self.total_words
        self.metadata["page_count"] = self.page_count
        self.metadata["file_type"] = self.file_type

    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "full_text": self.full_text,
            "pages": self.pages,
            "metadata": self.metadata,
            "total_chars": self.total_chars,
            "total_words": self.total_words,
            "page_count": self.page_count,
            "is_valid": self.is_valid,
            "error": self.error,
        }

    def __repr__(self) -> str:
        return (
            f"<ExtractedDocument filename='{self.filename}' type='{self.file_type}' "
            f"pages={self.page_count} words={self.total_words} valid={self.is_valid}>"
        )


class DocumentChunk:
    """
    Represents a discrete text chunk extracted from a source document.
    """
    def __init__(self, text: str, source: str, chunk_index: int, metadata: Optional[Dict[str, Any]] = None):
        self.text = text
        self.source = source
        self.chunk_index = chunk_index
        self.metadata = metadata or {}

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "source": self.source,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata,
        }


# -------------------------------------------------------------
# PDF TEXT EXTRACTION (PYPDF)
# -------------------------------------------------------------
def extract_text_from_pdf(
    file_source: Union[str, BinaryIO, bytes],
    filename: Optional[str] = None,
) -> ExtractedDocument:
    """
    Extracts text from a PDF file using pypdf, preserving per-page structure.
    
    Args:
        file_source: File path, file-like object, or raw bytes.
        filename: Optional filename override.
        
    Returns:
        ExtractedDocument instance.
    """
    try:
        import pypdf
    except ImportError:
        return ExtractedDocument(
            filename=filename or "document.pdf",
            file_type="pdf",
            error="pypdf library is not installed. Please install it via `pip install pypdf`.",
            is_valid=False,
        )

    doc_name = filename or (os.path.basename(file_source) if isinstance(file_source, str) else "document.pdf")
    pages_data: List[Dict[str, Any]] = []
    full_text_parts: List[str] = []

    try:
        if isinstance(file_source, (bytes, bytearray)):
            stream = io.BytesIO(file_source)
        elif isinstance(file_source, str):
            if not os.path.exists(file_source):
                return ExtractedDocument(
                    filename=doc_name,
                    file_type="pdf",
                    error=f"File not found: {file_source}",
                    is_valid=False,
                )
            stream = open(file_source, "rb")
        else:
            stream = file_source

        reader = pypdf.PdfReader(stream)

        # Check if encrypted/password protected
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return ExtractedDocument(
                    filename=doc_name,
                    file_type="pdf",
                    error="PDF is encrypted and password-protected.",
                    is_valid=False,
                )

        total_pages = len(reader.pages)
        if total_pages == 0:
            return ExtractedDocument(
                filename=doc_name,
                file_type="pdf",
                error="PDF contains zero pages or is empty.",
                is_valid=False,
            )

        for page_idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
                clean_page_text = page_text.strip()
                pages_data.append({
                    "page_number": page_idx,
                    "text": clean_page_text,
                    "char_count": len(clean_page_text),
                    "word_count": len(clean_page_text.split()) if clean_page_text else 0,
                })
                if clean_page_text:
                    full_text_parts.append(f"--- [Page {page_idx}] ---\n{clean_page_text}")
            except Exception as page_err:
                pages_data.append({
                    "page_number": page_idx,
                    "text": "",
                    "char_count": 0,
                    "word_count": 0,
                    "warning": f"Could not extract page {page_idx}: {str(page_err)}",
                })

        full_text = "\n\n".join(full_text_parts).strip()
        
        if not full_text:
            return ExtractedDocument(
                filename=doc_name,
                file_type="pdf",
                full_text="",
                pages=pages_data,
                error="No readable text found in PDF (file may be scanned images only or empty).",
                is_valid=False,
            )

        return ExtractedDocument(
            filename=doc_name,
            file_type="pdf",
            full_text=full_text,
            pages=pages_data,
            metadata={"total_pages": total_pages},
            is_valid=True,
        )

    except Exception as e:
        return ExtractedDocument(
            filename=doc_name,
            file_type="pdf",
            error=f"PDF extraction error: {str(e)}",
            is_valid=False,
        )
    finally:
        if isinstance(file_source, str) and 'stream' in locals() and hasattr(stream, 'close'):
            stream.close()


# -------------------------------------------------------------
# DOCX TEXT EXTRACTION (PYTHON-DOCX)
# -------------------------------------------------------------
def extract_text_from_docx(
    file_source: Union[str, BinaryIO, bytes],
    filename: Optional[str] = None,
) -> ExtractedDocument:
    """
    Extracts text from a DOCX file using python-docx, including paragraphs and tables.
    
    Args:
        file_source: File path, file-like object, or raw bytes.
        filename: Optional filename override.
        
    Returns:
        ExtractedDocument instance.
    """
    try:
        import docx
    except ImportError:
        return ExtractedDocument(
            filename=filename or "document.docx",
            file_type="docx",
            error="python-docx library is not installed. Please install it via `pip install python-docx`.",
            is_valid=False,
        )

    doc_name = filename or (os.path.basename(file_source) if isinstance(file_source, str) else "document.docx")

    try:
        if isinstance(file_source, (bytes, bytearray)):
            stream = io.BytesIO(file_source)
        elif isinstance(file_source, str):
            if not os.path.exists(file_source):
                return ExtractedDocument(
                    filename=doc_name,
                    file_type="docx",
                    error=f"File not found: {file_source}",
                    is_valid=False,
                )
            stream = open(file_source, "rb")
        else:
            stream = file_source

        doc = docx.Document(stream)
        
        extracted_sections: List[str] = []
        
        # 1. Paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                extracted_sections.append(text)
                
        # 2. Tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_str = " | ".join(c for c in row_cells if c)
                if row_str:
                    table_rows.append(row_str)
            if table_rows:
                extracted_sections.append("\n".join(table_rows))

        full_text = "\n\n".join(extracted_sections).strip()

        if not full_text:
            return ExtractedDocument(
                filename=doc_name,
                file_type="docx",
                full_text="",
                error="DOCX document contains no readable text or is empty.",
                is_valid=False,
            )

        return ExtractedDocument(
            filename=doc_name,
            file_type="docx",
            full_text=full_text,
            pages=[{
                "page_number": 1,
                "text": full_text,
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
            }],
            metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
            is_valid=True,
        )

    except Exception as e:
        return ExtractedDocument(
            filename=doc_name,
            file_type="docx",
            error=f"DOCX extraction error: {str(e)}",
            is_valid=False,
        )
    finally:
        if isinstance(file_source, str) and 'stream' in locals() and hasattr(stream, 'close'):
            stream.close()


# -------------------------------------------------------------
# TXT / PLAIN TEXT EXTRACTION (NATIVE PYTHON)
# -------------------------------------------------------------
def extract_text_from_txt(
    file_source: Union[str, BinaryIO, bytes],
    filename: Optional[str] = None,
) -> ExtractedDocument:
    """
    Extracts text from plain text files with robust multi-encoding fallback.
    
    Args:
        file_source: File path, file-like object, or raw bytes.
        filename: Optional filename override.
        
    Returns:
        ExtractedDocument instance.
    """
    doc_name = filename or (os.path.basename(file_source) if isinstance(file_source, str) else "document.txt")
    ext = os.path.splitext(doc_name)[1].lower().lstrip(".") or "txt"

    try:
        if isinstance(file_source, (bytes, bytearray)):
            raw_bytes = bytes(file_source)
        elif isinstance(file_source, str):
            if not os.path.exists(file_source):
                return ExtractedDocument(
                    filename=doc_name,
                    file_type=ext,
                    error=f"File not found: {file_source}",
                    is_valid=False,
                )
            with open(file_source, "rb") as f:
                raw_bytes = f.read()
        else:
            raw_bytes = file_source.read()
            if isinstance(raw_bytes, str):
                raw_bytes = raw_bytes.encode("utf-8")

        if not raw_bytes:
            return ExtractedDocument(
                filename=doc_name,
                file_type=ext,
                full_text="",
                error="Text file is empty (0 bytes).",
                is_valid=False,
            )

        # Multi-encoding decoding sequence
        text = None
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]:
            try:
                text = raw_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            text = raw_bytes.decode("utf-8", errors="ignore")

        clean_text = text.strip()
        if not clean_text:
            return ExtractedDocument(
                filename=doc_name,
                file_type=ext,
                full_text="",
                error="Text file contains only whitespace or is empty.",
                is_valid=False,
            )

        return ExtractedDocument(
            filename=doc_name,
            file_type=ext,
            full_text=clean_text,
            pages=[{
                "page_number": 1,
                "text": clean_text,
                "char_count": len(clean_text),
                "word_count": len(clean_text.split()),
            }],
            metadata={"byte_size": len(raw_bytes)},
            is_valid=True,
        )

    except Exception as e:
        return ExtractedDocument(
            filename=doc_name,
            file_type=ext,
            error=f"Text extraction error: {str(e)}",
            is_valid=False,
        )


# -------------------------------------------------------------
# UNIFIED DOCUMENT LOADER & DISPATCHER
# -------------------------------------------------------------
def load_and_extract_document(
    file_source: Union[str, BinaryIO, bytes],
    filename: Optional[str] = None,
    save_to_dir: Optional[str] = None,
) -> ExtractedDocument:
    """
    Unified loader that identifies file type and extracts text accordingly.
    Optionally saves the raw file to save_to_dir.
    
    Args:
        file_source: File path, file-like object, or bytes.
        filename: Name of the file with extension.
        save_to_dir: Optional directory path to persist the file.
        
    Returns:
        ExtractedDocument instance.
    """
    if filename is None:
        if isinstance(file_source, str):
            filename = os.path.basename(file_source)
        else:
            filename = getattr(file_source, "name", "document.txt")

    ext = os.path.splitext(filename)[1].lower()

    # Read raw bytes if saving is requested
    saved_path = None
    if save_to_dir:
        os.makedirs(save_to_dir, exist_ok=True)
        saved_path = os.path.join(save_to_dir, filename)
        try:
            if isinstance(file_source, (bytes, bytearray)):
                with open(saved_path, "wb") as f:
                    f.write(file_source)
            elif isinstance(file_source, str):
                if file_source != saved_path and os.path.exists(file_source):
                    shutil.copy2(file_source, saved_path)
            else:
                # Seek to 0, read bytes, seek back
                pos = file_source.tell() if hasattr(file_source, "tell") else 0
                data = file_source.read()
                with open(saved_path, "wb") as f:
                    f.write(data if isinstance(data, bytes) else data.encode("utf-8"))
                if hasattr(file_source, "seek"):
                    file_source.seek(pos)
        except Exception as save_err:
            pass

    if ext == ".pdf":
        doc = extract_text_from_pdf(file_source, filename=filename)
    elif ext == ".docx":
        doc = extract_text_from_docx(file_source, filename=filename)
    elif ext in [".txt", ".md", ".py", ".csv", ".json", ".rst", ".log", ".html"]:
        doc = extract_text_from_txt(file_source, filename=filename)
    else:
        doc = ExtractedDocument(
            filename=filename,
            file_type=ext.lstrip("."),
            error=f"Unsupported file format: '{ext}'. Supported formats: PDF, TXT, DOCX.",
            is_valid=False,
        )

    if saved_path:
        doc.metadata["saved_path"] = saved_path

    return doc


# -------------------------------------------------------------
# FILE SYSTEM & STORAGE HELPERS
# -------------------------------------------------------------
def get_documents_directory() -> str:
    """
    Returns the absolute path to the local documents directory.
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, "data", "documents")
    os.makedirs(docs_dir, exist_ok=True)
    return docs_dir


def save_uploaded_file(
    file_bytes: bytes,
    filename: str,
    target_dir: Optional[str] = None,
) -> str:
    """
    Saves raw file bytes into the target documents directory.
    
    Returns:
        Absolute path to the saved file.
    """
    if target_dir is None:
        target_dir = get_documents_directory()
    os.makedirs(target_dir, exist_ok=True)

    # Sanitize filename
    clean_name = os.path.basename(filename).replace("..", "").strip()
    if not clean_name:
        clean_name = f"document_{int(time.time())}.txt"

    file_path = os.path.join(target_dir, clean_name)
    with open(file_path, "wb") as f:
        f.write(file_bytes)

    return file_path


def list_saved_documents(directory_path: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Lists all saved documents in data/documents with size, date, and type.
    """
    if directory_path is None:
        directory_path = get_documents_directory()

    if not os.path.exists(directory_path):
        return []

    docs = []
    allowed_exts = {".pdf", ".docx", ".txt", ".md", ".py", ".csv", ".json"}

    for item in os.listdir(directory_path):
        full_path = os.path.join(directory_path, item)
        if os.path.isfile(full_path):
            ext = os.path.splitext(item)[1].lower()
            if ext in allowed_exts:
                stat = os.stat(full_path)
                size_kb = stat.st_size / 1024.0
                docs.append({
                    "filename": item,
                    "file_path": full_path,
                    "extension": ext.lstrip("."),
                    "size_bytes": stat.st_size,
                    "size_kb": size_kb,
                    "size_str": f"{size_kb:.1f} KB" if size_kb < 1024 else f"{size_kb/1024:.2f} MB",
                    "modified_at": stat.st_mtime,
                })

    # Sort newest first
    docs.sort(key=lambda d: d.get("modified_at", 0), reverse=True)
    return docs


def delete_saved_document(filename: str, directory_path: Optional[str] = None) -> bool:
    """
    Deletes a saved document from storage.
    """
    if directory_path is None:
        directory_path = get_documents_directory()

    clean_name = os.path.basename(filename)
    target_path = os.path.join(directory_path, clean_name)

    if os.path.exists(target_path) and os.path.isfile(target_path):
        try:
            os.remove(target_path)
            return True
        except Exception:
            return False
    return False


# -------------------------------------------------------------
# CHUNKING UTILITIES (FOR FUTURE STEPS)
# -------------------------------------------------------------
def chunk_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> List[str]:
    """
    Splits text into sliding-window character chunks with overlap.
    """
    if not text or not text.strip():
        return []

    cleaned = text.strip()
    if len(cleaned) <= chunk_size:
        return [cleaned]

    chunks = []
    start = 0
    step = max(1, chunk_size - chunk_overlap)

    while start < len(cleaned):
        end = min(len(cleaned), start + chunk_size)
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(cleaned):
            break
        start += step

    return chunks


def chunk_extracted_document(
    doc: ExtractedDocument,
    chunk_size: int = 500,
    chunk_overlap: int = 60,
) -> List[DocumentChunk]:
    """
    Chunks an ExtractedDocument while rigorously preserving page numbers and source metadata.
    
    Args:
        doc: ExtractedDocument instance.
        chunk_size: Target characters per chunk.
        chunk_overlap: Character overlap between consecutive chunks.
        
    Returns:
        List of DocumentChunk instances with rich metadata.
    """
    if not doc.is_valid or not doc.full_text.strip():
        return []

    chunks: List[DocumentChunk] = []
    global_idx = 0

    # 1. For PDFs with multiple pages, chunk per page to ensure exact page attribution
    if doc.file_type == "pdf" and doc.pages:
        for pg in doc.pages:
            pg_num = pg.get("page_number", 1)
            pg_text = pg.get("text", "").strip()
            if not pg_text:
                continue

            pg_chunks = chunk_text(pg_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            for chk_text in pg_chunks:
                chunks.append(
                    DocumentChunk(
                        text=chk_text,
                        source=doc.filename,
                        chunk_index=global_idx,
                        metadata={
                            "source": doc.filename,
                            "file_type": doc.file_type,
                            "page_number": pg_num,
                            "total_pages": doc.page_count,
                            "char_count": len(chk_text),
                            "word_count": len(chk_text.split()),
                        },
                    )
                )
                global_idx += 1
    else:
        # 2. For DOCX and TXT/MD, chunk full text
        text_chunks = chunk_text(doc.full_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        for chk_text in text_chunks:
            chunks.append(
                DocumentChunk(
                    text=chk_text,
                    source=doc.filename,
                    chunk_index=global_idx,
                    metadata={
                        "source": doc.filename,
                        "file_type": doc.file_type,
                        "page_number": 1,
                        "total_pages": doc.page_count or 1,
                        "char_count": len(chk_text),
                        "word_count": len(chk_text.split()),
                    },
                )
            )
            global_idx += 1

    return chunks


def load_single_file(
    file_path: str,
    chunk_size: int = 500,
    chunk_overlap: int = 60,
) -> List[DocumentChunk]:
    """
    Loads and chunks a single document file.
    """
    if not os.path.exists(file_path):
        return []

    doc = load_and_extract_document(file_path)
    return chunk_extracted_document(doc, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def load_and_chunk_all_documents(
    directory_path: Optional[str] = None,
    chunk_size: int = 500,
    chunk_overlap: int = 60,
) -> List[DocumentChunk]:
    """
    Loads and chunks all documents from the documents directory.
    """
    if directory_path is None:
        directory_path = get_documents_directory()

    if not os.path.exists(directory_path):
        return []

    all_chunks: List[DocumentChunk] = []
    saved_docs = list_saved_documents(directory_path)

    for doc_info in saved_docs:
        file_chunks = load_single_file(
            doc_info["file_path"],
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        all_chunks.extend(file_chunks)

    return all_chunks

